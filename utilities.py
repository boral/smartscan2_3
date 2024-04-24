import PIL.Image
import re
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Preformatted, Table, TableStyle, PageBreak
import fitz  # PyMuPDF
from reportlab.lib.pagesizes import letter, A4
import streamlit as st
import google.generativeai as genai
from dotenv import load_dotenv
import os
import time
import pandas as pd
from docx import Document
#from docx2pdf import convert
from io import BytesIO
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
import json
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx import Document
import subprocess

load_dotenv()

genai.configure( api_key=os.getenv('abc') )

model = genai.GenerativeModel('gemini-pro-vision')

image_prompt_0 = "You are an expert radiologist. Describe what is visible in given X ray or CT scan. Prepare very detailed and complete report as if given by a radiologist. Must contain Findings, Impressions, Recommendations and ICD-10 Code. In each case of Findings, Impressions and Recommendations give very detailed report as possible and add one or two line explaination also. You cannot tell that you are not capable of being a radiologist. You have to describe the image and don't give generic response like what a radiologist should do. Example output format: JSON with Findings, Recommendations, Impressions and ICD-10 Code as keys and each findings, each recommendations, each impressions and each ICD-10 code as sub values. Don't give triple quotes or triple backticks on top. Don't write the word json on top. Otherwise keep the remaining json formatting intact. Keep the [] for values of each keys."

image_prompt = 'Please provide a complete and detailed radiology report for the given X-ray or CT scan. The report should include the following sections: Findings, Impressions, Recommendations, ICD-10 Codes and Explanatory Notes. It is crucial to provide specific and detailed information for each section and avoid generic responses. Ensure that the headings for Findings, Impressions, Recommendations, ICD-10 codes and Explanatory Notes are clearly labeled.'

def login(username, password):
    credentials_df = pd.read_excel('credentials.xlsx')
    combined_credentials = username + '__' + password
    login_status = combined_credentials in list( credentials_df.combination )
    
    if login_status:
        role = list( credentials_df.loc[credentials_df['combination'] == combined_credentials, 'role'] )[0]
    else:
        role = None
    return login_status, role
    
def radiologist_report( input_image_path, image_prompt ):
    base64_image = PIL.Image.open(input_image_path)
    
    response = model.generate_content([ image_prompt, base64_image ], stream=True)
    response.resolve()
        
    response_text_1 = re.sub(r'```|\'\'\'', '', response.text)
        
    final_text = re.sub(r'\. ', '.\n', response_text_1)
        
    return final_text

def create_pdf(input_text, bottom_text, uploaded_image, patient_report_filename):
    try:
        # Convert iternation_num to a string if it is an integer
        bottom_text = str(bottom_text)
    
        # Create a PDF document
        doc = SimpleDocTemplate(patient_report_filename, pagesize=letter)
    
        # Create a list to hold the flowables (content elements) of the PDF
        flowables = []
    
        # Add a title to the document
        title_style = getSampleStyleSheet()["Title"]
        title = Paragraph("Report", title_style)
        flowables.append(title)
    
        # Add the input text to the document
        normal_style = getSampleStyleSheet()["Normal"]
        input_paragraph = Preformatted(input_text, normal_style)
        flowables.append(input_paragraph)
        
        # Add space between the main content and the bottom text
        flowables.append(Spacer(1, 20))
    
        # Add the image to the document
        if uploaded_image:
            image = Image(uploaded_image, width=300, height=300)  # Adjust width and height as needed
            flowables.append(image)
        
        # Add space between the main content and the bottom text
        flowables.append(Spacer(1, 20))
    
        # Add the supplied text message to the lower bottom
        bottom_style = getSampleStyleSheet()["Normal"]
        bottom_paragraph = Preformatted(bottom_text, bottom_style)
        flowables.append(bottom_paragraph)
    
        # Build the PDF document
        doc.build(flowables)
        
        return 'Success'
    except Exception as e:
        print(f"An error occurred: {e}")
        return 'Unsuccessful'
    
def create_docx(input_text, bottom_text, uploaded_image, patient_report_filename):
    try:
        # Create a new Word document
        doc = Document()

        # Add the input text to the document
        input_paragraph = doc.add_paragraph(input_text)

        # Add space between the main content and the bottom text
        doc.add_paragraph()

        # Add the supplied text message to the lower bottom
        doc.add_paragraph(bottom_text)

        # Save the Word document
        doc.save(patient_report_filename)

        return 'Success'
    except Exception as e:
        print(f"An error occurred: {e}")
        return 'Unsuccessful'
    
def display_pdf(pdf_path):

    cwd = os.getcwd()

    # Open the PDF file
    pdf_document = fitz.open(os.path.join(cwd, pdf_path))

    # Display each page of the PDF
    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)
        image_bytes = page.get_pixmap().tobytes()
        st.image(image_bytes, caption=f"Page {page_num + 1}", use_column_width=True)

    
def download_pdf( report_path, report_filename ):
    with open(report_path, "rb") as file:
        st.download_button(
                label="Download Report",
                data=file,
                file_name=report_filename
              )

def analysis( uploaded_image, patient_id, patient_image_filename ):
        
    patient_report_filename = 'generated_reports/report_' + patient_id + '.pdf'
    report_text = radiologist_report( patient_image_filename, image_prompt )
# =============================================================================
#     condition_met_flag = False
#     
#     words_to_look_for = ['Findings:', 'Impressions:', 'Recommendations:', 'Findings', 'Impressions', 'Recommendations' ]
#     
#     for iternation_num in range(5):
#         if any(word in report_text for word in words_to_look_for):
#             condition_met_flag = True
#             print("Condition met in iteration : ", str( iternation_num ) )
#             break
#         time.sleep(5)
#         report_text = radiologist_report( patient_image_filename, image_prompt )
#     
#     #... If even after retries, output is not obtained ...
#     if not condition_met_flag:
#         report_text = "Sorry ! Not able to process this image. Please try with some other image with better clarity."
# =============================================================================
    
    iternation_num = 0

    report_creation_flag = create_pdf( report_text, iternation_num, uploaded_image, patient_report_filename )   #... report
    
    return report_creation_flag


def analysis_docx( uploaded_image, patient_id, patient_image_filename ):
        
    patient_report_filename = 'generated_reports/report_' + patient_id + '.docx'
    report_text = radiologist_report( patient_image_filename, image_prompt )
    condition_met_flag = False
    
    words_to_look_for = ['Findings', 'Impressions', 'Recommendations' ]
    
    for iternation_num in range(5):
        if any(word in report_text for word in words_to_look_for):
            condition_met_flag = True
            print("Condition met in iteration : ", str( iternation_num ) )
            break
        time.sleep(5)
        report_text = radiologist_report( patient_image_filename, image_prompt )
    
    #... If even after retries, output is not obtained ...
    if not condition_met_flag:
        report_text = "Sorry ! Not able to process this image. Please try with some other image with better clarity."
    
    report_creation_flag = create_docx( report_text, iternation_num, uploaded_image, patient_report_filename )   #... report
    
    return report_creation_flag

def list_files(directory):
    cwd = os.getcwd()
    folder_path = os.path.join(cwd, directory)
    files = [file for file in os.listdir(folder_path) if os.path.isfile(os.path.join(folder_path, file))]
    file_list = ','.join(files)
    return file_list

def list_directories(directory):
    # Get the list of files and directories in the specified directory
    files_and_dirs = os.listdir(directory)
    
    # Filter out the directories from the list
    directories = [item for item in files_and_dirs if os.path.isdir(os.path.join(directory, item))]
    
    # Join the directories with commas
    directories_str = ', '.join(directories)
    
    # Return the comma-separated string
    return directories_str

def read_docx(file_path):
    doc = Document(file_path)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    return '\n'.join(text)

def write_docx_old(file, text):
    doc = Document()

    # Add text content to the document
    for line in text.split('\n'):
        doc.add_paragraph(line)

    doc.save(file)

def convert_docx_to_pdf(doc_path, path):

    subprocess.call(['soffice',
                 # '--headless',
                 '--convert-to',
                 'pdf',
                 '--outdir',
                 path,
                 doc_path])
    return doc_path

def remove_empty_pages(doc):
    for i in reversed(range(len(doc.sections))):
        section = doc.sections[i]
        for j in reversed(range(len(section.footer.paragraphs))):
            if not any(char.strip() for char in section.footer.paragraphs[j].text):
                section.footer.paragraphs.pop(j)
        for j in reversed(range(len(section.footer.tables))):
            if not any(cell.text.strip() for row in section.footer.tables[j].rows for cell in row.cells):
                section.footer.tables.pop(j)
    
def write_docx(text, image_path, df, docx_save_path, sign_path, logged_in_radiologist_details_df):
    doc = Document()

    # Add a dataframe table to the Word document
    table = doc.add_table(rows=len(df) + 1, cols=len(df.columns), style='Table Grid')
    table.autofit = True

    # Header row
    for col_num, col_name in enumerate(df.columns):
        cell = table.cell(0, col_num)
        cell.text = col_name
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)

    # Data rows
    for row_num, row_data in enumerate(df.values):
        for col_num, cell_value in enumerate(row_data):
            cell = table.cell(row_num + 1, col_num)
            cell.text = str(cell_value)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Add a line break
    doc.add_paragraph().space_after = Pt(0)
    
    # Add text content to the document
    for line in text.split('\n'):
        p = doc.add_paragraph(line)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
    
    doc.paragraphs[-1].clear()  # Remove the last paragraph, which is the empty one

    # Add radiologist sign
    doc.add_picture(sign_path, width=Pt(150), height=Pt(45))

    # Add radiologist details without extra spacing between lines
    # details_paragraphs = [ 'name', 'radiologist_degree', 'radiologist_designation', 'radiologist_registration_num' ]
    
    p0 = doc.add_paragraph( str( list(logged_in_radiologist_details_df['name'])[0] ) + ', ' + str( list(logged_in_radiologist_details_df['radiologist_degree'])[0] ) + ', ( ' + str( list(logged_in_radiologist_details_df['radiologist_designation'])[0] ) + ' )' )
    p0.paragraph_format.space_after = Pt(0)
    
    p = doc.add_paragraph( str( list(logged_in_radiologist_details_df['radiologist_registration_num'])[0]) )
    p.paragraph_format.space_after = Pt(0)
    
# =============================================================================
#     for detail_type in details_paragraphs:
#         detail_value = list(logged_in_radiologist_details_df[detail_type])[0]
#         p = doc.add_paragraph(detail_value)
#         p.paragraph_format.space_after = Pt(0)
# =============================================================================

    # Add a page break
    doc.add_page_break()

    # Add the uploaded image to the Word document
    doc.add_picture(image_path, width=Pt(400), height=Pt(500))
    
    # Remove empty pages
    remove_empty_pages(doc)

    doc.save(docx_save_path)


def save_as_docx_markdown(markdown_text, image_path, df, docx_save_path, sign_path, logged_in_radiologist_details_df):
    # Create a new Word document
    doc = Document()

    # Add a dataframe table to the Word document
    table = doc.add_table(rows=len(df) + 1, cols=len(df.columns), style='Table Grid')
    table.autofit = False

    # Header row
    for col_num, col_name in enumerate(df.columns):
        cell = table.cell(0, col_num)
        cell.text = col_name
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)

    # Data rows
    for row_num, row_data in enumerate(df.values):
        for col_num, cell_value in enumerate(row_data):
            cell = table.cell(row_num + 1, col_num)
            cell.text = str(cell_value)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Add a line break before the radiologist sign
    doc.add_paragraph()

    # Add radiologist sign
    doc.add_picture(sign_path, width=Pt(150), height=Pt(65))

    # Add radiologist details
    doc.add_paragraph(list(logged_in_radiologist_details_df['name'])[0] + ', ' + list(logged_in_radiologist_details_df['radiologist_degree'])[0] + ', ( ' + list(logged_in_radiologist_details_df['radiologist_designation'])[0] + ' )' )
    #doc.add_paragraph(list(logged_in_radiologist_details_df['radiologist_degree'])[0])
    #doc.add_paragraph(list(logged_in_radiologist_details_df['radiologist_designation'])[0])
    doc.add_paragraph(list(logged_in_radiologist_details_df['radiologist_registration_num'])[0])

    # Add a page break
    doc.add_page_break()

    # Add the uploaded image to the Word document
    doc.add_picture(image_path, width=Pt(400), height=Pt(500))

    # Save the Word document to the specified path
    doc.save(docx_save_path)
    # Create a new Word document
    doc = Document()

    # Add a dataframe table to the Word document
    table = doc.add_table(rows=len(df) + 1, cols=len(df.columns), style='Table Grid')
    table.autofit = False

    # Header row
    for col_num, col_name in enumerate(df.columns):
        cell = table.cell(0, col_num)
        cell.text = col_name
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)

    # Data rows
    for row_num, row_data in enumerate(df.values):
        for col_num, cell_value in enumerate(row_data):
            cell = table.cell(row_num + 1, col_num)
            cell.text = str(cell_value)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Add a line break before the radiologist sign
    doc.add_paragraph()

    # Add radiologist sign
    doc.add_picture(sign_path, width=Pt(150), height=Pt(65))

    # Add radiologist details
    doc.add_paragraph(list(logged_in_radiologist_details_df['name'])[0])
    doc.add_paragraph(list(logged_in_radiologist_details_df['radiologist_degree'])[0])
    doc.add_paragraph(list(logged_in_radiologist_details_df['radiologist_designation'])[0])
    doc.add_paragraph(list(logged_in_radiologist_details_df['radiologist_registration_num'])[0])

    # Add a page break
    doc.add_page_break()

    # Add the uploaded image to the Word document
    doc.add_picture(image_path, width=Pt(400), height=Pt(500))

    # Save the Word document to the specified path
    doc.save(docx_save_path)

def save_as_pdf_markdown(markdown_text, image_path, df, pdf_save_path, sign_path, logged_in_radiologist_details_df ):
    # Convert Markdown to HTML
    html_content = markdown_text # markdown2.markdown(markdown_text)

    # Create a BytesIO buffer to store the PDF content
    pdf_buffer = BytesIO()

    # Use ReportLab to generate PDF from HTML content
    pdf_doc = SimpleDocTemplate(pdf_buffer, pagesize=A4)  # Change pagesize to A4
    styles = getSampleStyleSheet()
    style = styles["BodyText"]

    # Add a dataframe to the PDF
    df_table = Table([df.columns] + df.values.tolist())
    df_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)  # Add this line for black borders
    ]))
    paragraphs = [Paragraph("", style), df_table]
    
    # Add a line break before the table
    paragraphs.append(Paragraph("", style))
    paragraphs.append(Paragraph("", style))

    # Add the Markdown content to the PDF using Preformatted
    markdown_paragraph = Preformatted(html_content, style)
    paragraphs.append(markdown_paragraph)
    
    # Add a line break before the radiologist sign
    paragraphs.append(Paragraph("", style))
    paragraphs.append(Paragraph("", style))
    
    # Add radiologist sign
    sign = Image(sign_path, hAlign = 'LEFT', width=150, height=65)
    paragraphs.append(sign)
    
    # Add radiologist details
    
    paragraphs.append(Paragraph( list( logged_in_radiologist_details_df['name'] )[0], style))
    paragraphs.append(Paragraph( list( logged_in_radiologist_details_df['radiologist_degree'] )[0], style))
    paragraphs.append(Paragraph( list( logged_in_radiologist_details_df['radiologist_designation'] )[0], style))
    paragraphs.append(Paragraph( list( logged_in_radiologist_details_df['radiologist_registration_num'] )[0], style))

    # Add a PageBreak to start a new page
    paragraphs.append(PageBreak())

    # Add the uploaded image to the PDF
    image = Image(image_path, width=400, height=500)
    paragraphs.append(image)

    # Build the PDF document
    pdf_doc.build(paragraphs)

    # Reset the buffer position to the beginning
    pdf_buffer.seek(0)

    # Save the PDF to the specified path
    with open(pdf_save_path, 'wb') as f:
        f.write(pdf_buffer.read())


def json_to_markdown(json_str):
    # Find the start and end of the JSON string within triple backticks
    start = json_str.find('{')
    end = json_str.rfind('}')

    # Extract the JSON part
    json_str = json_str[start:end+1].strip()

    try:
        # Try loading as JSON
        json_data = json.loads(json_str)
    except json.JSONDecodeError:
        raise ValueError("Invalid JSON syntax")

    markdown_content = ""

    for key, values in json_data.items():
        markdown_content += f"{key}\n"
        for value in values:
            markdown_content += f"- {value}\n"
        markdown_content += "\n"

    return markdown_content.strip()  # Remove leading and trailing whitespace



def save_as_pdf(markdown_content, output_pdf_path, headings, df):
    doc = SimpleDocTemplate(output_pdf_path, pagesize=letter)
    styles = getSampleStyleSheet()

    # Adjust fontSize for headings
    bold_heading_style = ParagraphStyle(
        'BoldHeading1',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=12  # Change font size to 12
    )
    styles.add(bold_heading_style)

    # Convert Markdown to PDF
    content = []

    # Add dataframe to content if provided
    df_table = Table([df.columns] + df.values.tolist())
    df_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)  # Add this line for black borders
    ]))
    content.append(Paragraph("", styles['Normal']))
    content.append(df_table)
    content.append(Spacer(1, 12))  # Add a blank line after the table

    # Add Markdown content to the PDF
    markdown_lines = markdown_content.split('\n')
    for line in markdown_lines:
        if line.strip():
            content.append(Paragraph(line, bold_heading_style if line.startswith(tuple(headings)) else styles['Normal']))

    doc.build(content)